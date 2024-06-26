from flask import *
from flask import jsonify #solo per errore too many request
from flask import render_template, send_file
from flask import flash
from flask import redirect, url_for, request, session, make_response
import shutil, os, datetime, random
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
from azure.core.credentials import AzureKeyCredential
from werkzeug.utils import secure_filename
import tempfile
import io 
from reportlab.pdfgen import canvas
import pdfplumber
from docx import Document
import docx
from flask import request
from flask_bcrypt import Bcrypt
from password_strength import PasswordPolicy
from zxcvbn import zxcvbn
import pwnedpasswords
from azure.ai.textanalytics import TextAnalyticsClient
from azure.identity import DefaultAzureCredential
from azure.core.exceptions import HttpResponseError
from dotenv import load_dotenv
import mysql.connector
from mysql.connector import Error
import re
import os
from argon2 import PasswordHasher
from io import BytesIO
from reportlab.lib.pagesizes import letter
import logging
from concurrent.futures import ThreadPoolExecutor

app = Flask(__name__, static_folder="static", template_folder="template")
app.secret_key = os.urandom(24)
bcrypt = Bcrypt(app)
connection_string = "DefaultEndpointsProtocol=https;AccountName=archiviocartelle;AccountKey=RSR+7NZLPvX8JCl2T/7zB29JcyC9lLe+BqRlFd63PeYP8urWUACSo1yrM2GiXibXi1QSyEReRo4m+AStJj6+ow==;EndpointSuffix=core.windows.net"
language_key = "ef9d0971353548fbaa27880a6cfc1039"
language_endpoint = "https://progettosistemi.cognitiveservices.azure.com/"
user="azure_admin_pii"
db_password="MjNORBNh$nbeKOyU"
blob_service_client = BlobServiceClient.from_connection_string('DefaultEndpointsProtocol=https;AccountName=archiviocartelle;AccountKey=RSR+7NZLPvX8JCl2T/7zB29JcyC9lLe+BqRlFd63PeYP8urWUACSo1yrM2GiXibXi1QSyEReRo4m+AStJj6+ow==;EndpointSuffix=core.windows.net')
policy = PasswordPolicy.from_names(length
                                   =12,)

#logging.basicConfig(level=logging.DEBUG)
def rf(a):
	return open(a,"r").read()

# Authenticate the client using your key and endpoint 
def authenticate_client():
    ta_credential = AzureKeyCredential(language_key)
    text_analytics_client = TextAnalyticsClient(
            endpoint=language_endpoint, 
            credential=ta_credential)
    return text_analytics_client

client = authenticate_client()
@app.after_request
def no_cache(response):
    """
    Add no caching headers.
    """
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    response.headers['Cache-Control'] = 'public, max-age=0'
    return response
# Connessione al database
def create_conn():
    conn = mysql.connector.connect(host='piiserver.mysql.database.azure.com',
                                   database='login_database',
                                   user=user,
                                   password=db_password)
    return conn

executor = ThreadPoolExecutor(max_workers=2)
def check_pwned_password(password):
    return pwnedpasswords.check(password)

@app.route('/')
def index():
    return render_template('index.html')

@app.route("/dashboard", methods=["GET"])
def dashboard():
    if 'username' in session:
        username = session['username']
        container_client = blob_service_client.get_container_client(username)
        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in container_client.list_blobs()]
        space = sum([blob.size for blob in container_client.list_blobs()]) / 1024 / 1024
        return render_template("dashboard.html", name=username, files=files, username=username, space=round(space), nf=len(files))
    else:
        return redirect(url_for('sign_in'))


def is_valid_username(username):
    return all(c.islower() or c.isdigit() for c in username)


@app.route("/sign_up", methods=["GET","POST"])
def sign_up():
    if request.method == 'POST':
        # Recupera i dati del modulo di registrazione
        password = request.form["password"]
        username = request.form["username"]
        
        # Verifica se la password è stata utilizzata in precedenza in modo asincrono
        future = executor.submit(check_pwned_password, password)

        # Verifica la lunghezza della password
        if policy.test(password):
            flash("Password must be at least 12 characters long", "error")
        print("Checking if password is too weak")
        # Verifica la forza della password
        if zxcvbn(password)['score'] < 3:
            flash("Password is too weak", "error")
        print("Checking if password has been pwned")
        # Verifica se la password è stata utilizzata in precedenza
        # Verifica se l'username è composto solo da lettere minuscole e numeri
        if not is_valid_username(username):
            print("Username must contain only lowercase letters and numbers")
            flash("Username must contain only lowercase letters and numbers", "error")
            return render_template("signup.html", message="Username must contain only lowercase letters and numbers")
        
        # Recupera i dati del modulo di registrazione
        user_data = {
            "username": request.form["username"],
            "password": bcrypt.generate_password_hash(request.form["password"]).decode('utf-8'),
            "email": request.form["email"],
            "name": request.form["name"]
        }
        # Controlla se l'utente o l'email esistono già
        try:
            print("Checking if user or email exists")
            with create_conn() as conn:
                with conn.cursor() as cursor:
                    cursor.execute("SELECT * FROM Users WHERE username = %s OR email = %s", (user_data["username"], user_data["email"]))
                    account = cursor.fetchone()
                    if (future.result()):                
                        return render_template("signup.html", message="Password has been pwned")
                    if account:
                        return render_template("signup.html", message="User or email already exists")
                    else: 
                        query = "INSERT INTO Users (username, password, email, nome_utente) VALUES (%s, %s, %s, %s)"
                        print(f"Executing query: {query}")
                        cursor.execute(query, (user_data["username"], user_data["password"], user_data["email"], user_data["name"]))
                        conn.commit()
        except Exception as e:
            print(f"Error checking if user or email exists: {e}")
            
        # Crea un contenitore per l'utente
        try:
            blob_service_client.create_container(user_data["username"])
            return render_template("dashboard.html", name=user_data["name"], files=[], password=user_data["password"], username=user_data["username"], space=0, nf=0)
            #return redirect(url_for('user_page', username=user_data['username']))
        except Exception as e:
            print(f"Error creating container for user: {e}")
            return render_template("signup.html", message="An error occurred: " + str(e))
    return render_template("signup.html", message="")
    
@app.route("/sign_in", methods=["GET","POST"])
def sign_in():
    if (request.method == "GET"):
        return render_template("signin.html", message="")
    else:
        username = request.form["username"]
        password = request.form["password"]
        try:
            conn = create_conn()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM Users WHERE username = %s", (username,))
            account = cursor.fetchone()
            if account:
                hashed_password = account[1]
                if bcrypt.check_password_hash(hashed_password, password):
                    # Set session and redirect to dashboard
                    session['username'] = username
                    return redirect(url_for('dashboard', username=username))
                else:
                    return render_template("signin.html", message="Wrong password")
            else:
                return render_template("signin.html", message="No such account")
        except Exception as e:
            print(f"Error checking if user exists: {e}")
            return render_template("signin.html", message="An error occurred")
        

@app.route("/upload/<user>/file", methods=["POST"])
def upload(user):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    if request.method == "POST":
        container_client = blob_service_client.get_container_client(user)
        files = request.files.getlist("files[]")
        valid_files = []

        # Validate files first
        for f in files:
            original_filename = f.filename
            name, extension = os.path.splitext(original_filename)
            if extension not in ['.docx', '.pdf']:
                flash("Invalid file type. Only .docx and .pdf files are allowed.")
                return redirect(url_for('dashboard', username=user))
            valid_files.append(f)

        # If validation passes, proceed with uploading
        for f in valid_files:
            original_filename = f.filename
            counter = 1
            file_blob = container_client.get_blob_client(f.filename)
            name, extension = os.path.splitext(f.filename)
            while file_blob.exists():
                new_filename = f"{name}({counter}){extension}"
                file_blob = container_client.get_blob_client(new_filename)
                counter += 1
            file_blob.upload_blob(f.read())

            document_name = file_blob.blob_name
            try:
                if extension == '.docx':
                    txt = docx_to_string(file_blob)
                elif extension == '.pdf':
                    txt = pdf_to_string(file_blob)
                documents = [txt]
                metadata, confidence_score = pii_detection(documents, user, extension, file_blob, container_client)
                metadata_json = json.dumps(metadata)
                conn = create_conn()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO documents (document_name, username, metadata) VALUES (%s, %s, %s)",
                    (document_name, user, metadata_json)
                )
                conn.commit()
                cursor.close()
                conn.close()
            except HttpResponseError as e:
                if e.status_code == 413:
                    flash("Document too large. Limit document size to: 5120 text elements.")
                    return redirect(url_for('dashboard', username=user,confidence_score=confidence_score))
                else:
                    raise

        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in container_client.list_blobs()]
        space = sum([blob.size for blob in container_client.list_blobs()]) / 1024 / 1024
        flash("Privacy Risk Level: " + confidence_score)
        return redirect(url_for('dashboard', username=user))
    else:
        return "Invalid request method", 405

@app.route("/convert/<user>/name/<name>", methods=["GET"])
def convert_text_to_anonimizedtext(user,name):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    if (request.method == "GET"):
            files_container = blob_service_client.get_container_client(user)
            file_blob = files_container.get_blob_client(name)
            if file_blob.exists():
                print("File exists")
                name,extension=os.path.splitext(name)
                if extension == '.docx':
                    txt= docx_to_string(file_blob)
                elif extension == '.pdf':
                    txt= pdf_to_string(file_blob)
                documents = [txt]
                output_file = "anonimazyed_" + name + extension

                # Eseguire il riconoscimento PII e redigere il testo
                language_country = client.detect_language(documents, country_hint="us")[0]
                language=language_country.primary_language['iso6391_name']
                print("LINGUA",language)
                response = client.recognize_pii_entities(documents, language=language)
                redacted_texts = []
                for doc in response:
                    if not doc.is_error:
                        redacted_texts.append(doc.redacted_text)
                    else:
                        print("Error:", doc.error.message)

                if redacted_texts:
                    #print("Converting text",redacted_texts)
                    if extension == ".pdf":
                        output = BytesIO()
                        c = canvas.Canvas(output, pagesize=letter)
                        width, height = letter
                        line_height = 14
                        margin = 100
                        y = height - margin
                        for text in redacted_texts:
                            lines = text.splitlines()
                            for line in lines:
                                if y < margin:
                                    c.showPage()
                                    y = height - margin
                                c.drawString(margin, y, line)
                                y -= line_height
                        c.save()
                        output.seek(0)
                        output_blob = files_container.get_blob_client(output_file)
                        output_blob.upload_blob(output.read(), overwrite=True)


                    elif extension == ".docx":
                        output = BytesIO()
                        doc = Document()
                        for text in redacted_texts:
                            doc.add_paragraph(text)
                        doc.save(output)
                        output.seek(0)
                        output_blob = files_container.get_blob_client(output_file)
                        output_blob.upload_blob(output.read(), overwrite=True)
                    
                    print("Testo convertito salvato come", output_file)
                else:
                    print("No PII detected")
                files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in files_container.list_blobs()]
                space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
                flash("Privacy Risk Level: No Risk")
                return redirect(url_for('dashboard', username=user))
    else:
            return "Invalid request method", 405

def docx_to_string(blob_client):
    blob_data = blob_client.download_blob().readall()
    document = Document(io.BytesIO(blob_data))
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])
    return text

def pdf_to_string(blob_client):
    blob_data = blob_client.download_blob().readall()
    with pdfplumber.open(io.BytesIO(blob_data)) as pdf:
        text = ' '.join([page.extract_text() for page in pdf.pages])
    return text
def pii_detection(documents, name, extension, file_blob, files_container):
    metadata_file = "metadata_" + name + ".json"
    response = client.recognize_pii_entities(documents, language="en")
    confidence_scores = []  # Aggiungi una lista per salvare i punteggi di confidenza
    metadata={}
    pii_found = False
    for doc in response:
            if not doc.is_error:
                for entity in doc.entities:
                    confidence_score = entity.confidence_score
                    confidence_scores.append(confidence_score)  # Salva il punteggio di confidenza
                    metadata[entity.text] = entity.category
            else:
                print("Error:", doc.error.message)
    if not pii_found:
        return {}, "No PII found"
    print("Metadata:", metadata)

    # Calcola il punteggio di confidenza medio e lo mostra a schermo
    if confidence_scores:
            # Estrai l'Average Confidence Score dai metadati, se esiste
            average_confidence_score = sum(confidence_scores) / len(confidence_scores)
            #Dice il livello di confidenza
            if average_confidence_score==0:
                text="No Risk"
            if average_confidence_score < 0.25:
                text="Very Low"
                return metadata, text
            elif average_confidence_score < 0.5:
                text="Low"
                return metadata, text
            elif average_confidence_score < 0.75:
                text="Medium"
                return metadata, text
            else:
                text="High"
                return metadata, text

@app.route("/delete/<user>/<name>", methods=["GET"])
def delete(user, name):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    print("Cancellazione file")
    if (request.method == "GET"):
                try:
                    files_container = blob_service_client.get_container_client(user)
                    file_blob = files_container.get_blob_client(name)
                    if file_blob.exists():
                        file_blob.delete_blob()
                        # If the file name does not start with "anonimized", delete it from the database
                        if not name.startswith("anonimized"):
                            try:
                                conn = create_conn()
                                cursor = conn.cursor()
                                cursor.execute("DELETE FROM documents WHERE document_name = %s AND username = %s", (name, user))
                                conn.commit()
                                cursor.close()
                                conn.close()
                            except Exception as e:
                                print(f"Error deleting document from database: {e}")
                                return "An error occurred during database deletion", 500
                        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in files_container.list_blobs()]
                        space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
                        return redirect(url_for('dashboard', username=user))
                    else:
                        return "File does not exist", 404
                except Exception as e:
                    print(f"Error during file deletion: {e}")
                    return "An error occurred during file deletion", 500
    return "Invalid request method", 405



@app.route("/download/<user>/name/<name>", methods=["GET"])
def download(user, name):
    if 'username' not in session:
        return redirect(url_for('index'))
    if (request.method == "GET"):
                files_container = blob_service_client.get_container_client(user)
                file_blob = files_container.get_blob_client(name)
                if file_blob.exists():
                    download_file_path = os.path.join(tempfile.gettempdir(), name)
                    with open(download_file_path, "wb") as download_file:
                        download_file.write(file_blob.download_blob().readall())
                    return send_from_directory(tempfile.gettempdir(), name, as_attachment=True)
    else:
            return "Invalid request method", 405
    
@app.route("/logout", methods=["GET"])
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.route('/search/<user>', methods=['POST','GET'])
def search(user):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    
    if request.method == 'POST':
        query = request.form['search']
        if query:
            conn = create_conn()
            cursor = conn.cursor(dictionary=True)
            
            # Perform search on document name and metadata
            search_query = """
                SELECT * FROM documents 
                WHERE username = %s AND 
                (document_name LIKE %s OR JSON_EXTRACT(metadata, '$') LIKE %s)
            """
            search_term = f"%{query}%"
            cursor.execute(search_query, (user, search_term, search_term))
            
            results = cursor.fetchall()
            cursor.close()
            conn.close()
            files_db = [{"name": result["document_name"]} for result in results]
            print("Files", files_db)
            
            files_container = blob_service_client.get_container_client(user)
            files_db_names = [file["name"] for file in files_db]  # Get a list of all file names in files_db
            files = []
            for blob in files_container.list_blobs():
                if blob.name in files_db_names:  # Only include the file if its name is in files_db_names
                    files.append({"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"})
                    print("Files", files)
            
            space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
            
            # Fetch user details for rendering the dashboard
            return render_template("dashboard.html", username=user, space=round(space), nf=len(files), files=files)
        else:
            return redirect(url_for('dashboard', username=user))
    else:  # GET request
        return render_template("dashboard.html", username=user, space=round(space), nf=len(files), files=files)

if __name__ == '__main__':
    app.run()
